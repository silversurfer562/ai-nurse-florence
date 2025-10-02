"""
PDF Generation Service for Patient Documents
Uses ReportLab to create professional, print-ready patient education materials
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    Table, TableStyle, ListFlowable, ListItem, Image
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfgen import canvas
from io import BytesIO
from datetime import datetime
from typing import List, Dict, Any, Optional
import os


class PatientDocumentPDF:
    """Base class for generating patient education PDFs"""

    # Brand colors
    PRIMARY_COLOR = HexColor('#2563EB')  # Blue
    SECONDARY_COLOR = HexColor('#10B981')  # Green
    WARNING_COLOR = HexColor('#EF4444')  # Red
    INFO_COLOR = HexColor('#F59E0B')  # Amber
    BACKGROUND_COLOR = HexColor('#F3F4F6')  # Light gray

    def __init__(self, title: str = "Patient Education Material"):
        self.title = title
        self.buffer = BytesIO()
        self.doc = SimpleDocTemplate(
            self.buffer,
            pagesize=letter,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch
        )
        self.styles = self._create_styles()
        self.story = []

    def _create_styles(self):
        """Create custom paragraph styles"""
        styles = getSampleStyleSheet()

        # Title style
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=self.PRIMARY_COLOR,
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))

        # Section header
        styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=self.PRIMARY_COLOR,
            spaceAfter=12,
            spaceBefore=18,
            fontName='Helvetica-Bold'
        ))

        # Sub-header
        styles.add(ParagraphStyle(
            name='SubHeader',
            parent=styles['Heading3'],
            fontSize=13,
            textColor=black,
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        ))

        # Body text
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=10
        ))

        # Warning box
        styles.add(ParagraphStyle(
            name='Warning',
            parent=styles['BodyText'],
            fontSize=11,
            leading=14,
            textColor=self.WARNING_COLOR,
            leftIndent=10,
            rightIndent=10,
            fontName='Helvetica-Bold'
        ))

        # Important note
        styles.add(ParagraphStyle(
            name='Important',
            parent=styles['BodyText'],
            fontSize=11,
            leading=14,
            textColor=self.INFO_COLOR,
            leftIndent=10,
            rightIndent=10,
            fontName='Helvetica-Bold'
        ))

        # Footer
        styles.add(ParagraphStyle(
            name='Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=HexColor('#6B7280'),
            alignment=TA_CENTER
        ))

        return styles

    def add_header(self, title: str, subtitle: Optional[str] = None):
        """Add document header"""
        self.story.append(Paragraph(title, self.styles['CustomTitle']))
        if subtitle:
            self.story.append(Paragraph(subtitle, self.styles['BodyText']))
        self.story.append(Spacer(1, 0.2 * inch))

    def add_section(self, title: str, content: str):
        """Add a section with title and content"""
        self.story.append(Paragraph(title, self.styles['SectionHeader']))
        self.story.append(Paragraph(content, self.styles['CustomBody']))
        self.story.append(Spacer(1, 0.15 * inch))

    def add_bullet_list(self, title: str, items: List[str], style_name: str = 'CustomBody'):
        """Add a bulleted list"""
        if not items:
            return

        self.story.append(Paragraph(title, self.styles['SectionHeader']))

        bullet_items = []
        for item in items:
            bullet_items.append(ListItem(
                Paragraph(item, self.styles[style_name]),
                leftIndent=20,
                bulletColor=self.PRIMARY_COLOR,
                value='circle'
            ))

        self.story.append(ListFlowable(
            bullet_items,
            bulletType='bullet',
            start='circle'
        ))
        self.story.append(Spacer(1, 0.15 * inch))

    def add_warning_box(self, title: str, items: List[str]):
        """Add a warning box with red border"""
        if not items:
            return

        data = [[Paragraph(f"<b>⚠️ {title}</b>", self.styles['Warning'])]]
        for item in items:
            data.append([Paragraph(f"• {item}", self.styles['CustomBody'])])

        table = Table(data, colWidths=[6.5 * inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.WARNING_COLOR),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('BACKGROUND', (0, 1), (-1, -1), HexColor('#FEE2E2')),
            ('BOX', (0, 0), (-1, -1), 2, self.WARNING_COLOR),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))

        self.story.append(table)
        self.story.append(Spacer(1, 0.2 * inch))

    def add_info_box(self, title: str, content: str):
        """Add an info box with blue border"""
        data = [
            [Paragraph(f"<b>ℹ️ {title}</b>", self.styles['Important'])],
            [Paragraph(content, self.styles['CustomBody'])]
        ]

        table = Table(data, colWidths=[6.5 * inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), self.INFO_COLOR),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('BACKGROUND', (0, 1), (-1, -1), HexColor('#FEF3C7')),
            ('BOX', (0, 0), (-1, -1), 2, self.INFO_COLOR),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))

        self.story.append(table)
        self.story.append(Spacer(1, 0.2 * inch))

    def add_medication_table(self, medications: List[Dict[str, str]]):
        """Add a table of medications"""
        if not medications:
            return

        self.story.append(Paragraph("Your Medications", self.styles['SectionHeader']))

        # Table header
        data = [['Medication', 'Dosage', 'Frequency', 'Instructions']]

        # Add medication rows
        for med in medications:
            data.append([
                med.get('name', ''),
                med.get('dosage', ''),
                med.get('frequency', ''),
                med.get('instructions', '')
            ])

        table = Table(data, colWidths=[2 * inch, 1.2 * inch, 1.5 * inch, 1.8 * inch])
        table.setStyle(TableStyle([
            # Header styling
            ('BACKGROUND', (0, 0), (-1, 0), self.PRIMARY_COLOR),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),

            # Body styling
            ('BACKGROUND', (0, 1), (-1, -1), white),
            ('TEXTCOLOR', (0, 1), (-1, -1), black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),

            # Grid
            ('GRID', (0, 0), (-1, -1), 1, HexColor('#E5E7EB')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),

            # Alternating row colors
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [white, HexColor('#F9FAFB')]),
        ]))

        self.story.append(table)
        self.story.append(Spacer(1, 0.2 * inch))

    def add_footer_disclaimer(self):
        """Add footer disclaimer"""
        self.story.append(Spacer(1, 0.3 * inch))

        disclaimer = (
            "<i>This document is for educational purposes only and does not replace professional medical advice. "
            "Always consult with your healthcare provider for medical questions or concerns.</i>"
        )
        self.story.append(Paragraph(disclaimer, self.styles['Footer']))

        generation_info = f"<i>Generated by AI Nurse Florence on {datetime.now().strftime('%B %d, %Y')}</i>"
        self.story.append(Paragraph(generation_info, self.styles['Footer']))

    def build_pdf(self) -> BytesIO:
        """Build the PDF and return as BytesIO"""
        self.doc.build(self.story, onFirstPage=self._add_page_number, onLaterPages=self._add_page_number)
        self.buffer.seek(0)
        return self.buffer

    def _add_page_number(self, canvas, doc):
        """Add page numbers to footer"""
        page_num = canvas.getPageNumber()
        text = f"Page {page_num}"
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        canvas.setFillColor(HexColor('#6B7280'))
        canvas.drawRightString(7.5 * inch, 0.5 * inch, text)
        canvas.restoreState()


class DischargeInstructionsPDF(PatientDocumentPDF):
    """Generate discharge instructions PDF"""

    def generate(self, data: Dict[str, Any]) -> BytesIO:
        """Generate discharge instructions PDF"""

        # Header
        title = "Discharge Instructions"
        if data.get('patient_name'):
            subtitle = f"For: {data['patient_name']}"
        else:
            subtitle = None

        self.add_header(title, subtitle)

        # Discharge date and diagnosis
        if data.get('discharge_date'):
            date_str = data['discharge_date'].strftime('%B %d, %Y')
            self.add_info_box("Discharge Date", date_str)

        if data.get('primary_diagnosis'):
            self.add_section("Reason for Visit / Diagnosis", data['primary_diagnosis'])

        # Medications
        if data.get('medications'):
            self.add_medication_table(data['medications'])

        # Follow-up appointments
        if data.get('follow_up_appointments'):
            self.add_bullet_list("Follow-Up Appointments", data['follow_up_appointments'])

        # Activity restrictions
        if data.get('activity_restrictions'):
            self.add_bullet_list("Activity Restrictions", data['activity_restrictions'])

        # Diet instructions
        if data.get('diet_instructions'):
            self.add_section("Diet Instructions", data['diet_instructions'])

        # Wound care
        if data.get('wound_care'):
            self.add_section("Wound Care", data['wound_care'])

        # Equipment
        if data.get('equipment_needs'):
            self.add_bullet_list("Equipment You May Need", data['equipment_needs'])

        # Warning signs
        if data.get('warning_signs'):
            self.add_warning_box("Warning Signs - Contact Your Doctor If You Experience:", data['warning_signs'])

        # Emergency criteria
        if data.get('emergency_criteria'):
            self.add_warning_box("Call 911 or Go to Emergency Room If:", data['emergency_criteria'])

        # Home care services
        if data.get('home_care_services'):
            self.add_info_box("Home Care Services", data['home_care_services'])

        # Footer
        self.add_footer_disclaimer()

        return self.build_pdf()


class MedicationGuidePDF(PatientDocumentPDF):
    """Generate medication guide PDF"""

    def generate(self, data: Dict[str, Any]) -> BytesIO:
        """Generate medication guide PDF"""

        medication_name = data.get('medication_name', 'Medication')

        # Header
        self.add_header(f"{medication_name} - Medication Guide")

        # Basic information
        info_text = f"<b>Dosage:</b> {data.get('dosage', 'As prescribed')}<br/>"
        info_text += f"<b>Frequency:</b> {data.get('frequency', 'As prescribed')}<br/>"
        info_text += f"<b>Route:</b> {data.get('route', 'Oral')}"
        self.add_section("Your Prescription", info_text)

        # Purpose
        if data.get('purpose'):
            self.add_section("Why You Are Taking This Medication", data['purpose'])

        # How it works
        if data.get('how_it_works'):
            self.add_section("How This Medication Works", data['how_it_works'])

        # Special instructions
        if data.get('special_instructions'):
            self.add_bullet_list("Important Instructions", data['special_instructions'])

        # Common side effects
        if data.get('common_side_effects'):
            self.add_bullet_list("Common Side Effects", data['common_side_effects'])

        # Serious side effects
        if data.get('serious_side_effects'):
            self.add_warning_box(
                "Serious Side Effects - Contact Your Doctor Immediately If You Experience:",
                data['serious_side_effects']
            )

        # Food interactions
        if data.get('food_interactions'):
            self.add_bullet_list("Foods to Avoid", data['food_interactions'])

        # Drug interactions
        if data.get('drug_interactions'):
            self.add_warning_box("Medications That May Interact", data['drug_interactions'])

        # Storage
        if data.get('storage_instructions'):
            self.add_info_box("Storage Instructions", data['storage_instructions'])

        # Missed dose
        if data.get('missed_dose_instructions'):
            self.add_info_box("If You Miss a Dose", data['missed_dose_instructions'])

        # Data sources
        if data.get('data_sources'):
            sources = ", ".join(data['data_sources'])
            self.add_section("Information Sources", f"This guide was generated using data from: {sources}")

        # Footer
        self.add_footer_disclaimer()

        return self.build_pdf()


class DiseaseEducationPDF(PatientDocumentPDF):
    """Generate disease education material PDF"""

    def generate(self, data: Dict[str, Any]) -> BytesIO:
        """Generate disease education PDF"""

        disease_name = data.get('disease_name', 'Medical Condition')

        # Header
        self.add_header(f"Understanding {disease_name}")

        # What it is
        if data.get('what_it_is'):
            self.add_section("What Is This Condition?", data['what_it_is'])

        # Causes
        if data.get('causes'):
            self.add_section("What Causes It?", data['causes'])

        # Symptoms
        if data.get('symptoms'):
            self.add_bullet_list("Common Signs and Symptoms", data['symptoms'])

        # Treatment options
        if data.get('treatment_options'):
            self.add_bullet_list("Treatment Options", data['treatment_options'])

        # Medications overview
        if data.get('medications_overview'):
            self.add_section("Medications", data['medications_overview'])

        # Self-care tips
        if data.get('self_care_tips'):
            self.add_bullet_list("Self-Care Tips", data['self_care_tips'])

        # Lifestyle modifications
        if data.get('lifestyle_modifications'):
            self.add_bullet_list("Lifestyle Changes That Can Help", data['lifestyle_modifications'])

        # Diet recommendations
        if data.get('diet_recommendations'):
            self.add_section("Diet and Nutrition", data['diet_recommendations'])

        # Exercise recommendations
        if data.get('exercise_recommendations'):
            self.add_section("Physical Activity", data['exercise_recommendations'])

        # Warning signs
        if data.get('warning_signs'):
            self.add_warning_box("Warning Signs - Contact Your Doctor If:", data['warning_signs'])

        # Emergency symptoms
        if data.get('emergency_symptoms'):
            self.add_warning_box("Call 911 If You Experience:", data['emergency_symptoms'])

        # Support groups
        if data.get('support_groups'):
            self.add_bullet_list("Support and Resources", data['support_groups'])

        # Additional resources
        if data.get('additional_resources'):
            self.add_bullet_list("Additional Resources", data['additional_resources'])

        # Questions to ask
        if data.get('questions_to_ask'):
            self.add_bullet_list("Questions to Ask Your Healthcare Provider", data['questions_to_ask'])

        # Data sources
        if data.get('data_sources'):
            sources = ", ".join(data['data_sources'])
            self.add_section("Information Sources", f"This guide was compiled using information from: {sources}")

        # Footer
        self.add_footer_disclaimer()

        return self.build_pdf()


def generate_discharge_instructions(data: Dict[str, Any]) -> BytesIO:
    """Helper function to generate discharge instructions PDF"""
    pdf = DischargeInstructionsPDF()
    return pdf.generate(data)


def generate_medication_guide(data: Dict[str, Any]) -> BytesIO:
    """Helper function to generate medication guide PDF"""
    pdf = MedicationGuidePDF()
    return pdf.generate(data)


def generate_disease_education(data: Dict[str, Any]) -> BytesIO:
    """Helper function to generate disease education PDF"""
    pdf = DiseaseEducationPDF()
    return pdf.generate(data)


# ============================================================================
# Word (DOCX) and Text Generation Functions
# ============================================================================

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_discharge_instructions_docx(data: Dict[str, Any]) -> BytesIO:
    """Generate discharge instructions as Word document"""
    doc = Document()

    # Title
    title = doc.add_heading('Discharge Instructions', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Patient info
    if data.get('patient_name'):
        p = doc.add_paragraph()
        p.add_run(f"Patient: {data['patient_name']}").bold = True

    if data.get('discharge_date'):
        p = doc.add_paragraph()
        p.add_run(f"Discharge Date: {data['discharge_date']}").bold = True

    doc.add_paragraph()

    # Primary Diagnosis
    doc.add_heading('Primary Diagnosis', 1)
    doc.add_paragraph(data.get('primary_diagnosis', 'Not specified'))

    # Medications
    if data.get('medications'):
        doc.add_heading('Medications', 1)
        for med in data['medications']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{med.get('name', '')} - {med.get('dosage', '')}").bold = True
            p.add_run(f"\n  Frequency: {med.get('frequency', '')}")
            if med.get('instructions'):
                p.add_run(f"\n  Instructions: {med.get('instructions', '')}")

    # Follow-up Appointments
    if data.get('follow_up_appointments'):
        doc.add_heading('Follow-Up Appointments', 1)
        for appt in data['follow_up_appointments']:
            doc.add_paragraph(appt, style='List Bullet')

    # Activity Restrictions
    if data.get('activity_restrictions'):
        doc.add_heading('Activity Restrictions', 1)
        for restriction in data['activity_restrictions']:
            doc.add_paragraph(restriction, style='List Bullet')

    # Diet Instructions
    if data.get('diet_instructions'):
        doc.add_heading('Diet Instructions', 1)
        doc.add_paragraph(data['diet_instructions'])

    # Warning Signs
    if data.get('warning_signs'):
        doc.add_heading('Warning Signs - Call Your Doctor If:', 1)
        for sign in data['warning_signs']:
            doc.add_paragraph(sign, style='List Bullet')

    # Emergency Criteria
    if data.get('emergency_criteria'):
        doc.add_heading('⚠️ CALL 911 or Go to ER Immediately If:', 1)
        for criteria in data['emergency_criteria']:
            doc.add_paragraph(criteria, style='List Bullet')

    # Wound Care
    if data.get('wound_care'):
        doc.add_heading('Wound Care Instructions', 1)
        doc.add_paragraph(data['wound_care'])

    # Equipment Needs
    if data.get('equipment_needs'):
        doc.add_heading('Equipment and Supplies', 1)
        for equipment in data['equipment_needs']:
            doc.add_paragraph(equipment, style='List Bullet')

    # Home Care Services
    if data.get('home_care_services'):
        doc.add_heading('Home Care Services', 1)
        doc.add_paragraph(data['home_care_services'])

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("This document is for informational purposes only and does not replace medical advice from your healthcare provider.").italic = True

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_discharge_instructions_text(data: Dict[str, Any]) -> BytesIO:
    """Generate discharge instructions as plain text"""
    lines = []
    lines.append("=" * 80)
    lines.append("DISCHARGE INSTRUCTIONS")
    lines.append("=" * 80)
    lines.append("")

    # Patient info
    if data.get('patient_name'):
        lines.append(f"Patient: {data['patient_name']}")
    if data.get('discharge_date'):
        lines.append(f"Discharge Date: {data['discharge_date']}")
    lines.append("")

    # Primary Diagnosis
    lines.append("PRIMARY DIAGNOSIS")
    lines.append("-" * 80)
    lines.append(data.get('primary_diagnosis', 'Not specified'))
    lines.append("")

    # Medications
    if data.get('medications'):
        lines.append("MEDICATIONS")
        lines.append("-" * 80)
        for med in data['medications']:
            lines.append(f"• {med.get('name', '')} - {med.get('dosage', '')}")
            lines.append(f"  Frequency: {med.get('frequency', '')}")
            if med.get('instructions'):
                lines.append(f"  Instructions: {med.get('instructions', '')}")
            lines.append("")

    # Follow-up Appointments
    if data.get('follow_up_appointments'):
        lines.append("FOLLOW-UP APPOINTMENTS")
        lines.append("-" * 80)
        for appt in data['follow_up_appointments']:
            lines.append(f"• {appt}")
        lines.append("")

    # Activity Restrictions
    if data.get('activity_restrictions'):
        lines.append("ACTIVITY RESTRICTIONS")
        lines.append("-" * 80)
        for restriction in data['activity_restrictions']:
            lines.append(f"• {restriction}")
        lines.append("")

    # Diet Instructions
    if data.get('diet_instructions'):
        lines.append("DIET INSTRUCTIONS")
        lines.append("-" * 80)
        lines.append(data['diet_instructions'])
        lines.append("")

    # Warning Signs
    if data.get('warning_signs'):
        lines.append("WARNING SIGNS - CALL YOUR DOCTOR IF:")
        lines.append("-" * 80)
        for sign in data['warning_signs']:
            lines.append(f"• {sign}")
        lines.append("")

    # Emergency Criteria
    if data.get('emergency_criteria'):
        lines.append("⚠️  CALL 911 OR GO TO ER IMMEDIATELY IF:")
        lines.append("-" * 80)
        for criteria in data['emergency_criteria']:
            lines.append(f"• {criteria}")
        lines.append("")

    # Wound Care
    if data.get('wound_care'):
        lines.append("WOUND CARE INSTRUCTIONS")
        lines.append("-" * 80)
        lines.append(data['wound_care'])
        lines.append("")

    # Equipment Needs
    if data.get('equipment_needs'):
        lines.append("EQUIPMENT AND SUPPLIES")
        lines.append("-" * 80)
        for equipment in data['equipment_needs']:
            lines.append(f"• {equipment}")
        lines.append("")

    # Home Care Services
    if data.get('home_care_services'):
        lines.append("HOME CARE SERVICES")
        lines.append("-" * 80)
        lines.append(data['home_care_services'])
        lines.append("")

    # Disclaimer
    lines.append("=" * 80)
    lines.append("This document is for informational purposes only and does not")
    lines.append("replace medical advice from your healthcare provider.")
    lines.append("=" * 80)

    # Join lines and encode to bytes
    text_content = "\n".join(lines)
    buffer = BytesIO(text_content.encode('utf-8'))
    buffer.seek(0)
    return buffer
